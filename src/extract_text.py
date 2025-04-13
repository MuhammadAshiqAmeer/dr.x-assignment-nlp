import os
from pathlib import Path
from docx import Document
import pandas as pd
import PyPDF2
import fitz  # PyMuPDF
from src.utils import save_text


def extract_text_from_file(file_path):
    """Extract text and figure captions from various file formats."""
    try:
        extension = os.path.splitext(file_path)[1].lower()
        text_output = []

        if extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text_output.append(f.read())

        elif extension == '.docx':
            doc = Document(file_path)
            for para in doc.paragraphs:
                text_output.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    text_output.append(row_text)

        elif extension == '.pdf':
            # Standard text extraction
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text_output.append(f"Page {page_num + 1}: {page.extract_text()}")

            # EXTRA: Figure caption extraction using PyMuPDF
            doc = fitz.open(file_path)
            for i, page in enumerate(doc):
                blocks = page.get_text("blocks")
                for block in blocks:
                    text = block[4].strip()
                    if text.lower().startswith("figure"):
                        caption = f"[Figure Caption - Page {i + 1}]: {text}"
                        text_output.append(caption)

        elif extension in ['.csv', '.xlsx', '.xls', '.xlsm']:
            if extension == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            text_output.append(df.to_string(index=False))

        else:
            raise ValueError(f"Unsupported file format: {extension}")

        extracted_text = '\n'.join([t for t in text_output if t.strip()])
        # Safely create an output path using the original filename
        relative_name = Path(file_path).stem + ".txt"
        output_dir = Path("outputs/extracted")
        output_dir.mkdir(parents=True, exist_ok=True)
        save_path = output_dir / relative_name

        save_text(extracted_text, str(save_path))
        return extracted_text

    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        return None

    

if __name__ == "__main__":
    base_dir = Path(__file__).parent.parent
    data_dir = os.path.join(base_dir,'data')
    
    if not os.path.exists(data_dir):
        print(f"Directory '{data_dir}' does not exist.")
    else:
        for root, _, files in os.walk(data_dir):
            for file in files:
                file_path = os.path.join(root, file)
                print(f"Processing: {file_path}")
                extracted = extract_text_from_file(file_path)
                if extracted:
                    print(f"Extraction successful: {file}")
                else:
                    print(f"Failed to extract text from: {file}")
